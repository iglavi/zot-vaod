"""חילוץ טקסט ומטא-דאטה מקבצי פסקי דין.

תומך גם ב-.docx (הפורמט האמיתי מאתר נט-המשפט) וגם ב-.txt (גיבוי/בדיקות).
מחלץ מגוף המסמך פרטים שאינם קיימים בקובץ ה-CSV: שם השופט/ת ותאריך ההחלטה.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from .config import HEB_MONTHS

_MONTHS_ALT = "|".join(HEB_MONTHS)

# "לפני כב' השופטת רוני סלע" / "לפני כבוד הרשם ..." וכד' — חלק מהמסמכים
# (כ-2,200 בפועל) משתמשים ב'בפני' במקום 'לפני'; בלעדיו judge יצא ריק
# לכולם. גם ':' אחרי לפני/בפני ('בפני: כבוד...') חייב להיות אופציונלי —
# אחרת ה-search מדלג על השורה הנכונה (הראשונה, עם השם) ותופס בטעות אזכור
# כללי מאוחר יותר בגוף ההחלטה (למשל 'יובא לפני כבוד הרשם' בלי שם). גם
# ספרה בודדת שנדבקת ל'לפני/בפני' (כ-176 מסמכים בפועל, כנראה שריד של
# הערת-שוליים/מספר עמוד שנשרך אל תוך זרימת הטקסט) חייבת להיות אופציונלית.
# חלון הלכידה (150 תווים) גדול מספיק כדי לכלול מותב של כמה שופטים —
# ראו _split_judge_panel, שמפצל את הקטע הזה למותב שלם, לא רק שופט/ת אחד/ת.
# re.DOTALL הכרחי: כשמותב מפורט בשורות נפרדות ('כבוד...\nכבוד...\nכבוד...',
# לא מודבק ברצף בלי שורה חדשה) — בלעדיו '.' לא חוצה \n, ולכידת הקבוצה
# הייתה נעצרת אחרי השופט/ת הראשון/ה בלבד, לפני שפיצול-המותב מקבל סיכוי
# בכלל לראות את השאר.
_JUDGE_RE = re.compile(r"[לב]פנ[יי]\s*\d*\s*:?\s*כב(?:ו?ד|['׳])?\s*(.{0,150})", re.DOTALL)
# תבנית תאריך לועזי בתוך גוף פסק הדין: "01 ינואר 2026"
_DATE_RE = re.compile(r"(\d{1,2})\s+(" + _MONTHS_ALT + r")\s+(\d{4})")
# גיבוי: שורת החתימה הסטנדרטית בסוף פסקי דין רבים — "ניתן/ניתנה היום,
# ט' בסיוון התשפ"ו (25.5.2026)." בלי זה, תאריך ההחלטה נשאר ריק במסמכים
# שבהם התאריך הלועזי מופיע רק שם (לא ליד 'לפני'/תחילת המסמך) — במיוחד
# בפסקי דין, לרוב בית ~19% ממסמכי הביניים עם decision_date ריק בפועל.
# ממוקד ל'ניתן/ניתנה' ליד תאריך בסוגריים — לא לתאריך מספרי גורף בטקסט
# (שעלול להיות תאריך הגשה/דיון אחר, לא תאריך ההחלטה עצמה) — ולוקח את
# ההתאמה *האחרונה* בטקסט, כי שורת החתימה כמעט תמיד האחרונה מסוגה.
_DATE_SIGNOFF_RE = re.compile(r"(?:ניתן|ניתנה)\b.{0,60}?\((\d{1,2})\.(\d{1,2})\.(\d{2,4})\)")
# מילים שמסמנות את סוף שם השופט/ת (תארים, תפקידי צדדים, מונחי פתיח).
# הוסר 'נ\'' (קיצור 'נגד'): מתנגש עם ראשי-תיבות של שמות שופטים שמתחילים
# באות נ' (למשל 'נ\' הנדל') וגוזר את שם המשפחה. 'נגד' המלא, שאינו מעורפל
# באותו אופן, עדיין תופס.
_JUDGE_STOP = re.compile(
    r"העתק|פסק[\s\-]?דין|החלט|בעניין|בין\b|נגד|מיום|כבוד|"
    r"מבקש|משיב|עורר|מערער|תוב[ ע]|נתבע|עות[ר]|המבקש|בקשה|רקע|"
    r"ת[\"״]?א|\d|\n"
)
# פיצול מותב של כמה שופטים שהודבקו זה לזה בלי רווח/שורה חדשה בחילוץ מה-PDF
# ('...עמיתכבוד השופט...') — 'כבוד' עצמו מסמן תחילת שופט/ת נוסף/ת במותב.
_JUDGE_PANEL_SPLIT = re.compile(r"כבוד")


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _read_docx(p: Path) -> str:
    """קורא טקסט מלא מ-.docx ישירות מה-XML הגולמי (לא דרך python-docx
    document.paragraphs/tables). הכרחי כי תבניות המסמכים הרשמיים של בתי
    המשפט משתמשות ב'content controls' (שדות מילוי: שם השופט/ת, מספר
    התיק, הצדדים) שעוטפים את הפסקה כולה ב-<w:sdt><w:sdtContent>. פסקה
    כזו אינה צאצא ישיר של גוף המסמך, ולכן document.paragraphs מדלג עליה
    לגמרי — מה שגרם לשדות מטא-דאטה קריטיים לצאת ריקים כמעט תמיד."""
    import zipfile
    from lxml import etree

    with zipfile.ZipFile(str(p)) as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)
    paragraphs = []
    for p_el in root.iter(f"{{{_W_NS}}}p"):
        text = "".join(t.text or "" for t in p_el.findall(f".//{{{_W_NS}}}t"))
        if text.strip():
            paragraphs.append(text)
    return "\n".join(paragraphs)


# ארכיון ישן (קבצים עם סיומות לא-סטנדרטיות כמו .B01/.N02/.V01, כנראה
# מהמרה/סריקה ישנה בת עשרות שנים) שבו טקסט עברי נשמר על הדיסק כשהוא כבר
# פגום מלכתחילה: בתים בקידוד CP1255 (עברית) פוענחו בטעות כ-CP1252
# ('ANSI' מערבי) ונשמרו מחדש — כך ש-'בית המשפט' נהיה 'áéú äîùôè'. זה
# UTF-8 תקין לגמרי (אין שגיאת קידוד לתפוס), רק שהתוכן שגוי מיסודו —
# ולכן כל שדות המטא-דאטה (חסרי תווים עבריים אמיתיים להתאמה) יוצאים ריקים
# עבור מסמכים כאלה. מתקן בהמרה הפוכה: קידוד בחזרה כ-CP1252 ופענוח כ-CP1255
# משחזר את הטקסט העברי המקורי במדויק. 'כבוד' (לא 'בית המשפט') משמש סימן-
# היכר: מופיע כמעט בכל מסמך (שורת הצגת השופט/ים), בעוד ש'בית המשפט'
# נכתב לפעמים עם מקף ('בבית-המשפט') ולא רווח — מפספס חלק מהמסמכים.
_MOJIBAKE_SIGNATURE = "ëáåã"  # 'כבוד' אחרי CP1255->CP1252


def _fix_cp1255_mojibake(text: str) -> str:
    if not text or _MOJIBAKE_SIGNATURE not in text:
        return text
    try:
        return text.encode("cp1252").decode("cp1255")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return text


# תווי בקרת-כיווניות בלתי-נראים (LRM/RLM ואח') ש-Word מוסיף סביב מספרים/
# תאריכים בטקסט דו-כיווני — נעלמים בתצוגה אבל נדבקים ישירות בין '(' לספרה
# הראשונה בשורת החתימה ("ניתנה היום... (‏15.5.2017)"), ושוברים את
# _DATE_SIGNOFF_RE כי היא דורשת ספרה מיד אחרי הסוגר. הוסר ~90% מהמקרים
# שבהם decision_date נשאר ריק. אין להם שום תפקיד סמנטי (רק רמז-רינדור),
# כך שהסרה גורפת בטוחה.
_BIDI_MARKS_RE = re.compile(r"[‎‏‪-‮⁦-⁩]")


# ארכיון בית המשפט העליון (documents_supreme/) מכיל קבצים עם סיומת
# '.docx' שרובם ככולם (~99% בפועל, מדגם) אינם OOXML תקין בכלל — אלא
# קובצי RTF (~85%) או Word בינארי ישן מסוג OLE2 (~14%) ששמם שונה בטעות
# בזמן ארגון הארכיון (לפני עשרות שנים, לפני הפרויקט הזה). python-docx
# נכשל בשקט על שניהם, וה-fallback ל-PDF (שסובל מבעיות היפוך-כיוון וזיהוי
# תווים) הוא המקור לרוב בעיות הגיבריש/היפוך-טקסט שנמצאו הלילה. לכן
# מזהים את הסוג האמיתי לפי בתי הפתיחה של הקובץ (לא לפי הסיומת שלו) —
# RTF הוא טקסט נקי בלי הסיכון של PDF, כך שמקדימים אותו כשמזוהה.
_OLE2_SIGNATURE = bytes([0xD0, 0xCF, 0x11, 0xE0, 0xA1, 0xB1, 0x1A, 0xE1])


def _sniff_kind(p: Path) -> str:
    """מזהה את הסוג האמיתי של הקובץ לפי בתי הפתיחה, לא לפי סיומת השם.
    מחזיר 'docx' (ZIP/OOXML תקין), 'rtf', 'ole2' (Word בינארי ישן), או ''."""
    try:
        with p.open("rb") as f:
            header = f.read(8)
    except OSError:
        return ""
    if header[:2] == b"PK":
        return "docx"
    if header[:5] == b"{\\rtf":
        return "rtf"
    if header == _OLE2_SIGNATURE:
        return "ole2"
    return ""


def _read_rtf(p: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = p.read_text(encoding="cp1255", errors="replace")
    return rtf_to_text(raw)


# Word בינארי ישן (OLE2, ~14% מארכיון בית המשפט העליון) — python-docx לא
# תומך בפורמט הזה בכלל. antiword (כלי CLI ישן אך יציב, מגיע מובנה עם Git
# for Windows תחת mingw64/bin) יודע לקרוא אותו; '-m UTF-8.txt' חיוני כדי
# שהפלט יהיה UTF-8 תקין במקום CP862/לטיני שגוי כברירת המחדל.
def _read_ole2(p: Path) -> str:
    import subprocess

    try:
        result = subprocess.run(
            ["antiword", "-m", "UTF-8.txt", str(p)],
            capture_output=True,
            timeout=15,
        )
    except (OSError, subprocess.TimeoutExpired):
        return ""
    if result.returncode != 0:
        return ""
    text = result.stdout.decode("utf-8", errors="replace")
    # antiword משרטט טבלאות Word (כאן: שדות הפתיח - בית משפט/מספר
    # תיק/שופט/צדדים, כמעט תמיד בטבלה בתבניות הישנות) עם '|' כגבול עמודה.
    # ה-'|' נדבק ישירות לתוכן התא הבא בלי רווח אמיתי ביניהם, ושובר את
    # regex-י החילוץ שמצפים למילה הבאה מיד אחרי הכותרת (למשל 'לפני:|כבוד'
    # במקום 'לפני: כבוד') — בפרט judge, שכמעט תמיד יצא ריק בלעדי זה.
    return text.replace("|", " ")


def read_text(path: str | Path) -> str:
    """קורא טקסט מלא מקובץ pdf/docx/txt. מחזיר מחרוזת ריקה במקרה כשל."""
    p = Path(path)
    suffix = p.suffix.lower()
    try:
        if suffix in (".docx", ".doc"):
            kind = _sniff_kind(p)
            if kind == "rtf":
                text = _read_rtf(p)
            elif kind == "ole2":
                text = _read_ole2(p)
            elif suffix == ".docx":
                text = _read_docx(p)
            else:
                import docx  # python-docx (לא תומך רשמית ב-.doc הישן, best-effort)

                document = docx.Document(str(p))
                parts = [para.text for para in document.paragraphs]
                text = "\n".join(t for t in parts if t and t.strip())
        elif suffix == ".pdf":
            text = _read_pdf(p)
        else:
            text = p.read_text(encoding="utf-8", errors="ignore")
        return _BIDI_MARKS_RE.sub("", _fix_cp1255_mojibake(text))
    except Exception:
        return ""


# PDF-ים עבריים ישנים (בעיקר טפסי בתי-משפט מלפני ~2015) נוצרו בפונטים
# 'ויזואליים' שאינם תומכים ב-BiDi תקני — pdfplumber/pypdf מחלצים מהם את
# התווים לפי סדר הציור על הדף (משמאל לימין), כך שכל שורה עברית יוצאת
# הפוכה לגמרי. מזהים את התופעה לפי מילת-מפתח נפוצה שמופיעה הפוכה (ולא
# במצורה הרגילה) בטקסט שחולץ, ומתקנים עם python-bidi (base_dir='R'
# משמר נכון גם רצפי מספרים/אנגלית מוטמעים, בניגוד להיפוך שורה גולמי).
# 'בית המשפט' (עם ה"א הידיעה) לא תמיד מופיע — מסמכים רבים כותבים 'בית
# משפט השלום/העבודה' וכד' בלי ה"א על 'משפט' עצמו — ולכן גם 'לפני/בפני'
# ההפוך (שורת הצגת השופט/ים, כמעט אוניברסלי בכל מסמך) משמש סימן-היכר
# נוסף, כדי לא לפספס מסמכים הפוכים שלא כוללים את הביטוי המדויק הראשון.
_REVERSED_SIGNATURE = "טפשמה תיב"  # 'בית המשפט' הפוך
_REVERSED_SIGNATURE_2 = re.compile(r"ינפ[בל]")  # 'לפני'/'בפני' הפוך
_NORMAL_SIGNATURE = "בית המשפט"
_NORMAL_SIGNATURE_2 = re.compile(r"[לב]פני")


def _fix_visual_hebrew(text: str) -> str:
    if not text:
        return text
    has_reversed = _REVERSED_SIGNATURE in text or _REVERSED_SIGNATURE_2.search(text)
    has_normal = _NORMAL_SIGNATURE in text or _NORMAL_SIGNATURE_2.search(text)
    if not has_reversed or has_normal:
        return text
    try:
        from bidi.algorithm import get_display
    except ImportError:
        return text
    return "\n".join(get_display(line, base_dir="R") for line in text.split("\n"))


def _read_pdf(p: Path) -> str:
    """מחלץ טקסט מקובץ PDF. מנסה pdfplumber (טוב יותר לעברית) ואז pypdf."""
    try:
        import pdfplumber

        with pdfplumber.open(str(p)) as pdf:
            pages = [(page.extract_text() or "") for page in pdf.pages]
        text = "\n".join(pages).strip()
        if text:
            return _fix_visual_hebrew(text)
    except Exception:
        pass
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        return _fix_visual_hebrew(text)
    except Exception:
        return ""


def extract_judge(text: str) -> str:
    """מחלץ את שם/שמות השופט/ת (כולל תואר) משורת 'לפני כב\'...'.

    תומך במותב של כמה שופטים שהודבקו זה לזה בלי מעבר שורה בחילוץ מה-PDF
    ('...עמיתכבוד השופט...') — מפצל לפי 'כבוד' ומחזיר את כולם, מופרדים
    בפסיק, במקום רק את הראשון/ה."""
    if not text:
        return ""
    m = _JUDGE_RE.search(text)
    if not m:
        return ""
    seg = m.group(1)
    names = []
    for chunk in _JUDGE_PANEL_SPLIT.split(seg):
        piece = _JUDGE_STOP.split(chunk)[0]
        piece = re.sub(r"\s+", " ", piece).strip(" .,:-‏‎")
        if len(piece) >= 4:
            names.append(piece)
        else:
            break
    return ", ".join(names)


# טווח שנים סביר לבדיקת תקינות: לפני קום המדינה אין פסקי דין ישראליים,
# ומעבר לשנה הבאה מ'עכשיו' זו כמעט תמיד ספרה שסורקה/הוזחה בטעות (למשל
# תאריך מעורבב ממש"מ ישן: 'ניתן היום ... )13.32.39(' — חודש 32 לא קיים).
_MIN_PLAUSIBLE_YEAR = 1948
_MAX_PLAUSIBLE_YEAR = 2027


def _valid_iso_date(day: int, month: int, year: int) -> str:
    """מחזיר 'YYYY-MM-DD' רק אם זה תאריך אמיתי וסביר, אחרת ''."""
    if not (_MIN_PLAUSIBLE_YEAR <= year <= _MAX_PLAUSIBLE_YEAR):
        return ""
    try:
        date(year, month, day)  # מוודא יום/חודש חוקיים (כולל שנים מעוברות)
    except ValueError:
        return ""
    return f"{year:04d}-{month:02d}-{day:02d}"


def extract_decision_date(text: str) -> str:
    """מחלץ את תאריך ההחלטה (הלועזי) ומחזיר ISO 'YYYY-MM-DD', או ''."""
    if not text:
        return ""
    m = _DATE_RE.search(text)
    if m:
        result = _valid_iso_date(int(m.group(1)), HEB_MONTHS[m.group(2)], int(m.group(3)))
        if result:
            return result
    # גיבוי: שורת חתימה "ניתן/ניתנה היום... (DD.MM.YYYY)" — לוקחים את
    # ההתאמה האחרונה (החתימה כמעט תמיד אחרונה, לא הפניה לפסק דין אחר).
    for m2 in reversed(list(_DATE_SIGNOFF_RE.finditer(text))):
        day, month, year = int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
        if year < 100:
            # שנה דו-ספרתית: המאגר כולל פסקי דין מתחילת שנות ה-90 ואילך —
            # "99" הוא 1999 (לא 2099). חלון סביר: 00-30 -> 20XX, 31-99 -> 19XX.
            year += 2000 if year <= 30 else 1900
        result = _valid_iso_date(day, month, year)
        if result:
            return result
    return ""


# זיהוי כותרת פסק הדין: {בית משפט} {סוג תיק} {מספר תיק} {צדדים} לפני ...
_HEAD_RE = re.compile(
    r"(?P<ctype>[֐-׿\"״]{2,6})\s+"
    r"(?P<num>\d{2,6}[-/]\d{1,2}(?:-\d{2,4})?)"
)
# סימני-עצירה אמינים תמיד (פותחים בפועל את שורת השופט/ת או 'תיק חיצוני').
_PARTIES_STOP = re.compile(r"[לב]פנ[יי]|תיק\s+חיצוני")
# 'כבוד'/"כב'" כשלעצמו הוא סימון-עצירה חלש יותר: ארגונים בשם 'דרך כבוד'
# וכד' הכילו את המילה בשם הצד עצמו, וגרמו לקטיעה שגויה של השם. משמש רק
# כגיבוי אם לא נמצא 'לפני/בפני' תקין באותה שורה כלל.
_PARTIES_STOP_FALLBACK = re.compile(r"כב(?:ו?ד|['׳])")

# גיבוי לחילוץ צדדים: חלק מהמסמכים (בעיקר בג"ץ/החלטות ביניים) מציגים
# קודם את השופט/ים ('לפני: כבוד...') ורק אחר-כך, בתוויות מפורשות
# ('העותרים:'/'המבקשת:'/'המשיבים:' וכד'), את שמות הצדדים — הפוך מהסדר
# שה-extraction הרגיל מניח (צדדים ואז שופט/ים). בלי הגיבוי הזה parties
# יוצא ריק לגמרי, כי _PARTIES_STOP תופס את 'לפני' כמעט מיד אחרי מספר
# התיק. גם תומך בתיקים מאוחדים עם כמה תוויות עוקבות ('העותרים בבג"ץ
# X:\nהעותרים בבג"ץ Y:') — לוקח את השם הראשון מתחת לתווית הראשונה
# (התיק שמספרו כבר חולץ), לפי בקשת המשתמש: "אם יש כמה שמות בצד, שם
# ההליך הוא השם הראשון בכל צד".
# מאשימה/נאשם(ים): תוויות תיקים פליליים (בג"ץ/אזרחי משתמשים בתוויות
# האחרות) — בלעדיהן, כל תיק פלילי במבנה 'לפני שופט, ואז תוויות צד' נשאר
# עם parties ריק, כי אין לו כלל תובע/עותר/מבקש. אלה נכתבות בד"כ בלי ':'
# ('מאשימה\nמדינת ישראל') — לכן תבנית נפרדת ומצומצמת (חייבת שורה חדשה
# מיד אחרי המילה) ולא סתם הפיכת ':' לאופציונלי בתבנית הכללית: זה גרם
# להתאמות שווא במקומות אחרים בטקסט (למשל 'נגד' עם רווחים, שברי מספרי
# תיק) שלא היו תוויות-צד אמיתיות בכלל.
_PLAINTIFF_LABEL_RE = re.compile(
    r"(?:(?:העותר(?:ת|ים)?|המבקש(?:ת|ים)?|התובע(?:ת|ים)?|המערער(?:ת|ים)?)[^:\n]{0,40}:|"
    r"(?:מאשימה|התביעה)\s*\n)"
)
_DEFENDANT_LABEL_RE = re.compile(
    r"(?:(?:המשיב(?:ה|ים)?|הנתבע(?:ת|ים)?)[^:\n]{0,40}:|"
    r"(?:נאשמ(?:ים|ת)?|הנאשמ(?:ים|ת)?)\s*\n)"
)
# חלק-משנה של תיק מאוחד, בכתיב גימטרייה ("- ל", "- כ\"א", "- י\"ד") —
# בדיוק המקרה שהערת ה-TODO למעלה (ליד "תיקים מאוחדים... חלק-משנה כאות
# בודדת") מזהה, אבל בפועל לא סינן: אות בודדת עם גרש ("ד'") או שתי אותיות
# עם גרשיים ("כ"ז") עוברות את סינון האורך (2 <= len <= 80) כי הגרש/גרשיים
# מוסיפים תו, כך שהתוצאה נשמרה כאילו הייתה שם צד אמיתי. בלט הרבה יותר
# בקבצי OLE2/antiword (שם השורה "ע"א 1356/09 - כ"א" מודפסת כשורה אחת
# בטבלה, בלי ':' שיפריד אותה מהשם), אבל התבנית קיימת גם במקורות אחרים.
_GEMATRIA_ORDINAL_RE = re.compile(r"^[א-ת](?:['׳]|[\"״][א-ת])$")
_VS_RE = re.compile(r"\bנגד\b")
_LIST_ITEM_START_RE = re.compile(r"^\s*\d+[.)]\s*")
_PARTY_NAME_STOP_RE = re.compile(
    r"\n|\d+[.)]|" + _PLAINTIFF_LABEL_RE.pattern + "|" + _DEFENDANT_LABEL_RE.pattern + r"|\bנגד\b"
)


def _skip_stacked_labels(segment: str, label_re: re.Pattern) -> str:
    """מדלג על תוויות-צד עוקבות ('העותרים בבג"ץ X:\\nהעותרים בבג"ץ Y:')
    עד לתחילת התוכן בפועל — כדי שלא ניחשב את התווית השנייה כ'שם הצד'."""
    s = segment
    while True:
        stripped = s.lstrip()
        m = label_re.match(stripped)
        if not m:
            return stripped
        s = stripped[m.end():]


def _first_party_name(segment: str) -> tuple[str, bool]:
    """מחזיר (השם הראשון, האם יש עוד שמות באותו צד) מתוך הקטע שבין
    תווית הצד לעצירה הבאה (שורה חדשה, פריט ממוספר הבא, תווית אחרת, 'נגד')."""
    s = segment.lstrip()
    s = _LIST_ITEM_START_RE.sub("", s, count=1)
    m = _PARTY_NAME_STOP_RE.search(s)
    name = (s[:m.start()] if m else s[:80]).strip(" .,:-–'\"׳״")
    if not m:
        has_more = False
    elif re.match(r"\d+[.)]", m.group()):
        has_more = True  # העצירה עצמה היא הפריט הממוספר הבא
    else:
        has_more = bool(re.match(r"\d+[.)]", s[m.end():].lstrip()))
    return name, has_more


def _extract_parties_after_judges(text: str) -> str:
    """גיבוי: מחפש תוויות צד מפורשות בתחילת המסמך (אחרי אזור השופט/ים),
    ובונה מהן 'שם ראשון נ' שם ראשון' (עם 'ואח'' אם יש עוד באותו צד)."""
    window = text[:2500]
    p_label = _PLAINTIFF_LABEL_RE.search(window)
    if not p_label:
        return ""
    vs = _VS_RE.search(window, p_label.end())
    d_label = _DEFENDANT_LABEL_RE.search(window, vs.end() if vs else p_label.end())
    if not d_label:
        return ""
    p_content = _skip_stacked_labels(window[p_label.end():(vs.start() if vs else d_label.start())], _PLAINTIFF_LABEL_RE)
    d_content = _skip_stacked_labels(window[d_label.end():], _DEFENDANT_LABEL_RE)
    p_name, p_more = _first_party_name(p_content)
    d_name, d_more = _first_party_name(d_content)
    if not (2 <= len(p_name) <= 80) or not (2 <= len(d_name) <= 80):
        return ""
    p_disp = p_name + (" ואח'" if p_more else "")
    d_disp = d_name + (" ואח'" if d_more else "")
    return f"{p_disp} נ' {d_disp}"


# התאמה נחשבת כותרת אמיתית רק אם נמצאה קרוב לתחילת המסמך. אחרת, כנראה
# שההתאמה מקרית (איזשהו מספר בטקסט הגוף, לא מספר התיק בפועל) — וניפול
# למקרה 'לא נמצאה כותרת' (שדות ריקים) במקום למלא אותם בזבל.
_HEAD_MAX_POS = 150

# מסמכים מסוג 'מספר בקשה:N' (החלטות ביניים) לא כוללים בטקסט הגלוי שם
# בית משפט/סוג תיק/מספר תיק בפורמט הרגיל — רק את מספר הבקשה עצמה. בלי
# הזיהוי הזה, _HEAD_RE היה מוצא התאמה מקרית מאוחר יותר בטקסט ומזהם את
# כל השדות (כולל 'בית משפט') בגוש טקסט לא רלוונטי.
_BAKASHA_RE = re.compile(r"^מספר\s+בקשה\s*:?\s*\d+")

# כמעט כל סוג הליך אמיתי כולל מרכאות (ת"א, בג"ץ, חדל"פ...) — בלעדי
# הדרישה הזו, ה-ctype (2-6 תווים עבריים חופשי) תפס במקרה כל מילה קצרה
# שמופיעה במקרה ממש לפני תבנית 'מספר/מספר' כלשהי (תאריך בגוף ההחלטה,
# לרוב) — 'ליום'/'מיום'/'ביום' (תאריך), 'אזרחי'/'פלילי' (תיאור הליך),
# ואף שברי טקסט הפוך שלא תוקנו. זוהה בפועל שהתבנית הבלתי-מוגבלת יצרה
# עשרות ערכי-זבל שונים ב-case_type ובגררה אחריה גם court/case_number
# שגויים (כל הטקסט שלפני ה'התאמה' המקרית).
def _valid_ctype(ctype: str) -> bool:
    return ctype == "העליון" or bool(re.search(r"[\"'׳״]", ctype))


def _find_head_match(head: str):
    for m in _HEAD_RE.finditer(head):
        if m.start() > _HEAD_MAX_POS:
            return None
        if _valid_ctype(m.group("ctype")):
            return m
    return None


_COURT_START_RE = re.compile(r"בית\s*[-–]?\s*ה?משפט|ה?משפט\s+העליון|בג[\"'׳״]?ץ")
_SUPREME_START_RE = re.compile(r"^(בית\s*[-–]?\s*)?ה?משפט\s+העליון")
_BAGATZ_RE = re.compile(r"בג[\"'׳״]?ץ")
# שורות-זבל נפוצות שאינן שם בית משפט כלל (שם שופט/ת שנגרר, כותרות מסמך
# כלליות) — כשלא זוהתה תבנית 'בית משפט' תקינה, אלה היחידות שחוסמות את
# הערך המקורי; כל דבר אחר (למשל 'אזורי לעבודה חיפה', 'המחוזי ירושלים' —
# שמות מקוצרים של סוגי בתי דין/משפט שלא כוללים את המילה 'משפט' עצמה)
# נשמר כפי שהוא, כדי לא לאבד ערכי סינון תקינים-אך-חלקיים.
_JUNK_COURT_RE = re.compile(
    r"^(לפני|בפני)\b|^(תוכן עניינים|תאריך|תקציר|תמצית|מספר|המערער(ת)?|"
    r"תיקים מאוחדים|הודעה לתקשורת)\b"
)
# 'מחוזי חיפה' ו'בית המשפט המחוזי בחיפה' (וכנ"ל 'שלום X'/'בית משפט השלום
# בX') הם אותו בית משפט בדיוק — הצורה הקצרה מגיעה מ-CSV המקור (החלטות
# בית המשפט, לא ארכיון העליון) שמקצר לא-עקבית, ובלי איחוד הן נספרות
# ומוצגות כשני בתי משפט שונים במסנן.
_SHORT_COURT_RE = re.compile(r"^ה?(מחוזי|שלום)\s+(.+)$")
_SHORT_COURT_PREFIX = {"מחוזי": "בית המשפט המחוזי ב", "שלום": "בית משפט השלום ב"}
# 'בית משפט לתעבורה' הוא בית משפט נפרד, לא תת-מסלול של בית משפט השלום —
# חלק מהמסמכים כותבים בטעות 'בית משפט השלום לתעבורה', מה שיוצר גם כפילות
# מול הצורה הנכונה ('בית המשפט לתעבורה מחוז X').
_SHALOM_TRAFFIC_RE = re.compile(r"(בית\s*[-–]?\s*ה?משפט)\s+השלום(\s+לתעבורה)")
# שרידים שהחילוץ המקורי (_HEAD_RE, שרירת מספר-תיק גזלה את 'העליון' לתוך
# case_type — ראו extract_metadata) או טעות הקלדה בודדת במסמך המקור
# משאירים מאחור — בהקשר הזה חד-משמעית בית המשפט העליון (אין בית משפט
# ישראלי אחר בשם 'העליון' סתם, וזו טעות דפוס ברורה של 'בית המשפט העליון').
_SUPREME_EXACT_ALIASES = {"העליון", "בית המשפ העליון", "בבית המשפ העליון"}
# 'בית המשפט קמא' הוא ביטוי-הפניה גנרי לערכאה קודמת בהליך ערעור ('הערכאה
# הראשונה/הקודמת') — לא שם בית משפט ספציפי, ולעיתים נגרר אחריו טקסט
# ההפניה המלא (שם שופט/מספר תיק של הערכאה ההיא). לא ניתן לדעת מכך איזה
# בית משפט זה בפועל.
_KAMA_RE = re.compile(r"^בית\s*[-–]?\s*ה?משפט\s+קמא\b")
# תווי סוגריים שנשמרו הפוכים-כיוון (RTL) בחילוץ מ-PDF ('...)בת-ים(' במקום
# '(בת-ים)') — הכיוון הלוגי לא ידוע, אז פשוט מסירים את כל תווי הסוגריים
# ולא מנסים לשחזר איזה מהם פותח/סוגר.
_BRACKETS_RE = re.compile(r"[()\[\]{}]")


def _normalize_court(raw: str) -> str:
    """מנקה ומאחד את שם בית המשפט שחולץ מכותרת המסמך.

    מטפל בשלוש תקלות שנצפו בפועל בשדה הזה: (1) טקסט הפוך מ-PDF ישן —
    כמו ב-full_text (ראו _fix_visual_hebrew), רק שכאן המחרוזת קצרה מדי
    כדי שסימן ההיכר של _fix_visual_hebrew תמיד יופיע בה, ולכן הבדיקה
    כאן מסתמכת על היעדר 'משפט' בכיוון תקין; (2) תווי-זבל שנגררו לפני שם
    בית המשפט (סימני פיסוק, מספרי עמוד, שארית פסקה קודמת) — נחתכים על
    ידי איתור נקודת ההתחלה האמיתית של השם; (3) עשרות גרסאות שונות של
    'בית המשפט העליון' (עם/בלי 'בירושלים', 'תקציר', 'הודעה לתקשורת'
    וכו') שהופכות מסנן בית-משפט באתר לרשימה ארוכה ומבלבלת — לכל אלה יש
    רק בית משפט עליון אחד במדינה, ולכן מאוחדות לערך קנוני יחיד.
    מחזיר מחרוזת ריקה אם לא זוהה שם בית משפט תקין בכלל."""
    s = (raw or "").strip()
    if not s:
        return ""
    if s in _SUPREME_EXACT_ALIASES:
        return "בית המשפט העליון"
    if "משפט" not in s and not _BAGATZ_RE.search(s):
        try:
            from bidi.algorithm import get_display
            fixed = get_display(s, base_dir="R")
        except ImportError:
            fixed = s
        if "משפט" in fixed or _BAGATZ_RE.search(fixed):
            s = fixed
    m_short = _SHORT_COURT_RE.match(s)
    if m_short:
        kind, city = m_short.groups()
        if kind == "מחוזי" and city == "מרכז":
            # יוצא-דופן: שם המחוז ('מרכז') אינו שם עיר שמקבל 'ב' — הכינוי
            # המלא הנהוג הוא 'מחוז מרכז (לוד)', לא 'ב-מרכז'.
            s = "בית המשפט המחוזי מרכז לוד"
        else:
            s = _SHORT_COURT_PREFIX[kind] + city
    # לוקחים את ההתאמה האחרונה, לא הראשונה: לפעמים קודמת לכותרת האמיתית
    # שורת-הקדמה שמזכירה בית משפט אחר בהקשר שונה לגמרי (למשל הודעת איסור
    # פרסום 'לפי החלטת בית משפט מחוזי' לפני הכותרת האמיתית 'בבית המשפט
    # העליון') — ההתאמה הקרובה ביותר למספר התיק היא כמעט תמיד הנכונה.
    matches = list(_COURT_START_RE.finditer(s))
    if matches:
        s = s[matches[-1].start():]
        if _SUPREME_START_RE.match(s):
            return "בית המשפט העליון"
    elif _JUNK_COURT_RE.search(s) or len(s) > 40 or re.search(r"\d", s):
        return ""
    s = _SHALOM_TRAFFIC_RE.sub(r"\1\2", s)
    s = _BRACKETS_RE.sub("", s)
    s = re.sub(r"\s*[-–]\s*", " ", s)
    s = re.sub(r"\s+בהליך\s*$", "", s)
    s = re.sub(r"\s+", " ", s).strip(" '\"׳״;:.,]})|+")
    if s in ("בית המשפט", "בית משפט") or _KAMA_RE.match(s):
        # 'בית המשפט' לבדו (בלי שם/עיר/סוג), או 'בית משפט קמא' (הפניה
        # גנרית לערכאה קודמת, לא שם ספציפי) — אי אפשר לדעת מהם איזה בית
        # משפט זה בפועל.
        return ""
    return s if len(s) <= 40 and not re.search(r"\d", s) else ""


def extract_metadata(text: str) -> dict:
    """מחלץ מטא-דאטה מגוף פסק הדין (עבור קבצים ללא שורת CSV).

    מחזיר: court, case_type, case_number, parties, judge, decision_type,
    decision_date — כל שדה best-effort, מחרוזת ריקה אם לא נמצא."""
    out = {"court": "", "case_type": "", "case_number": "", "parties": "",
           "judge": "", "decision_type": "", "decision_date": ""}
    if not text:
        return out
    head = re.sub(r"\s+", " ", text.strip())[:600]

    m = None if _BAKASHA_RE.match(head) else _find_head_match(head)
    if m:
        ctype = m.group("ctype").strip()
        # 'העליון' הוא היוצא-מן-הכלל היחיד שכשיר כ-ctype בלי מרכאות (ראו
        # _valid_ctype) — כשהוא מופיע ממש לפני מספר תיק (כותרות רשם/משפט
        # ביניים מסוג 'בבית המשפט העליון 8382/07') זה חלק משם בית המשפט,
        # לא סוג ההליך, ולכן מצטרף לקטע ה-court במקום להיחשב ctype.
        if ctype == "העליון":
            court = head[:m.end("ctype")].strip(" -–:")
            ctype = ""
        else:
            court = head[:m.start()].strip(" -–:")
        out["court"] = _normalize_court(court)
        out["case_type"] = ctype
        out["case_number"] = m.group("num").strip()
        rest = head[m.end():]
        # תיקים מאוחדים/צמודים חושפים מספר-תיק שני (לפעמים גם שלישי) לפני
        # השורה עם 'לפני/בפני' — בלי לזהות אותו כסימן-עצירה נוסף, הוא היה
        # נגרר לתוך parties (או, אם ארוך מדי, גורם לדחיית השדה כולו).
        candidates = [s for s in (_PARTIES_STOP.search(rest), _find_head_match(rest)) if s]
        stop = min(candidates, key=lambda s: s.start()) if candidates else _PARTIES_STOP_FALLBACK.search(rest)
        parties = (rest[:stop.start()] if stop else rest).strip(" -–:‏‎")
        # מסננים "תיק חיצוני: 123/2025" אם נגרר
        parties = re.sub(r"תיק\s+חיצוני.*$", "", parties).strip(" -–:")
        # תיקים מאוחדים לפעמים מציינים חלק-משנה כאות/ספרה בודדת אחרי מספר
        # התיק ("בג"ץ 10016/16 – ל", "מ"ח 7929/96 - 1") — כשאין תוכן אמיתי
        # לפני 'לפני/בפני', זה מה שהיה נשאר ונשמר בטעות כאילו הוא שם הצד.
        if 2 <= len(parties) <= 80 and not _GEMATRIA_ORDINAL_RE.match(parties):
            out["parties"] = parties

    if not out["parties"]:
        out["parties"] = _extract_parties_after_judges(text)

    out["judge"] = extract_judge(text)
    out["decision_date"] = extract_decision_date(text)
    # הסדר כאן קריטי: מסמכי הכרעת-דין/גזר-דין (בעיקר תיקים פליליים) כמעט
    # תמיד כוללים גם את הביטוי הגנרי 'העתק החלטה מפרוטוקול' (תבנית-עטיפה
    # של המסמך) *לפני* הכותרת הספציפית שלהם — כך שבדיקת 'החלט' לבדה הייתה
    # מסווגת את כולם בטעות כ'החלטה' גנרית, בלי קשר לכך שהם דווקא הכרעת
    # דין/גזר דין (נבדק בפועל: ~38,000 מסמכים בקורפוס נפגעו מכך). לכן
    # הביטויים הספציפיים האלה נבדקים *לפני* הגיבוי הגנרי, לא רק 'פסק דין'.
    if "פסק דין" in head or "פסק-דין" in head:
        out["decision_type"] = "פסק דין"
    elif "גזר דין" in head or "גזר-דין" in head:
        out["decision_type"] = "גזר דין"
    elif "הכרעת דין" in head or "הכרעת-דין" in head:
        out["decision_type"] = "הכרעת דין"
    elif "החלט" in head:
        out["decision_type"] = "החלטה"
    return out


def filed_date_from_case(case_number: str) -> str:
    """גוזר תאריך פתיחה משוער ממספר התיק (למשל 49000-12-25 -> 2025-12-01)."""
    if not case_number:
        return ""
    parts = str(case_number).strip().split("-")
    if len(parts) == 3 and parts[1].isdigit() and parts[2].isdigit():
        mm, yy = int(parts[1]), parts[2]
        if 1 <= mm <= 12 and len(yy) == 2:
            return f"20{yy}-{mm:02d}-01"
    return ""
